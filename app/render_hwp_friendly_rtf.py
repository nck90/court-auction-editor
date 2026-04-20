#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from render_final_notice import GROUP_ORDER, format_entry, load_entries, parse_case_sort_key


TABLE_COLUMNS = [
    ("사건번호", 28),
    ("물건번호", 14),
    ("소재지 및 면적 [㎡]", 86),
    ("용도", 20),
    ("감정평가액 / 최저매각가격", 32),
    ("비고", 40),
]


def rtf_escape(text: str) -> str:
    out: list[str] = []
    for ch in text:
        code = ord(ch)
        if ch == "\\":
            out.append(r"\\")
        elif ch == "{":
            out.append(r"\{")
        elif ch == "}":
            out.append(r"\}")
        elif ch == "\n":
            out.append(r"\line ")
        elif 32 <= code < 127:
            out.append(ch)
        else:
            if code > 32767:
                code -= 65536
            out.append(rf"\u{code}?")
    return "".join(out)


def grouped_entries(doc: dict) -> dict[str, list[dict]]:
    entries = [format_entry(entry) for entry in doc["entries"]]
    grouped = {group: [] for group in GROUP_ORDER}
    for entry in entries:
        grouped[entry["group"]].append(entry)
    for group in GROUP_ORDER:
        grouped[group] = sorted(grouped[group], key=lambda row: parse_case_sort_key(row["case"], row["item"]))
    return grouped


def rtf_cell(text: str) -> str:
    return r"\intbl " + rtf_escape(text or "-") + r"\cell "


def render_rtf(doc: dict) -> str:
    grouped = grouped_entries(doc)
    widths = [1600, 800, 4800, 1200, 1800, 2200]
    cell_rights = []
    total = 0
    for width in widths:
        total += width
        cell_rights.append(total)

    parts = [
        r"{\rtf1\ansi\deff0",
        r"{\fonttbl{\f0 Apple SD Gothic Neo;}{\f1 Malgun Gothic;}}",
        r"\viewkind4\uc1\pard\lang1042\f0\fs20",
        r"\qc\b\fs28 " + rtf_escape("법원경매공고 한글 작업본") + r"\b0\fs20\par",
        r"\qc " + rtf_escape(doc.get("court_line") or "법원 경매부동산의 매각 공고") + r"\par",
        r"\ql " + rtf_escape(doc.get("auction_datetime") or "-") + r"\par",
        r"\ql " + rtf_escape(doc.get("decision_datetime") or "-") + r"\par",
        r"\ql " + rtf_escape(doc.get("officer_line") or "-") + r"\par\par",
    ]

    for group in GROUP_ORDER:
        rows = grouped[group]
        if not rows:
            continue
        parts.append(r"\pard\b\fs24 " + rtf_escape(f"[{group}]") + r"\b0\fs20\par")
        header = [name for name, _ in TABLE_COLUMNS]
        parts.append(r"\trowd\trgaph60")
        for right in cell_rights:
            parts.append(rf"\cellx{right}")
        for value in header:
            parts.append(rtf_cell(value))
        parts.append(r"\row")
        for row in rows:
            parts.append(r"\trowd\trgaph60")
            for right in cell_rights:
                parts.append(rf"\cellx{right}")
            values = [row["case"], row["item"], row["location"], row["usage"], row["price"], row["note"] or "-"]
            for value in values:
                parts.append(rtf_cell(value))
            parts.append(r"\row")
        parts.append(r"\par")
    parts.append("}")
    return "".join(parts)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_mm: int) -> None:
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_mm * 56.7)))
    tc_w.set(qn("w:type"), "dxa")


def write_cell(cell, text: str, *, bold: bool = False, center: bool = False, font_size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text or "-")
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if center else WD_ALIGN_VERTICAL.TOP


def render_docx(doc: dict, out_path: Path) -> None:
    grouped = grouped_entries(doc)
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("법원경매공고 한글 작업본")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(16)

    for line in [
        doc.get("court_line") or "법원 경매부동산의 매각 공고",
        doc.get("auction_datetime") or "-",
        doc.get("decision_datetime") or "-",
        doc.get("officer_line") or "-",
    ]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(9)

    for group in GROUP_ORDER:
        rows = grouped[group]
        if not rows:
            continue

        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(6)
        heading_run = heading.add_run(f"[{group}]")
        heading_run.bold = True
        heading_run.font.name = "Malgun Gothic"
        heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(31, 56, 100)

        table = document.add_table(rows=1, cols=len(TABLE_COLUMNS))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = False

        header_cells = table.rows[0].cells
        for idx, (title, width) in enumerate(TABLE_COLUMNS):
            set_cell_width(header_cells[idx], width)
            set_cell_shading(header_cells[idx], "D9E2F3")
            write_cell(header_cells[idx], title, bold=True, center=True, font_size=8)

        for row in rows:
            values = [row["case"], row["item"], row["location"], row["usage"], row["price"], row["note"] or "-"]
            tr = table.add_row().cells
            for idx, value in enumerate(values):
                set_cell_width(tr[idx], TABLE_COLUMNS[idx][1])
                write_cell(tr[idx], value, center=idx in {1, 3, 4}, font_size=8)

        document.add_paragraph()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="한글에서 바로 수정 가능한 DOCX/RTF 작업본을 생성합니다.")
    parser.add_argument("json_path", type=Path)
    parser.add_argument("-o", "--output-dir", type=Path, default=Path("output"))
    args = parser.parse_args()

    doc = load_entries(args.json_path)
    args.output_dir.mkdir(parents=True, exist_ok=True)
    rtf_path = args.output_dir / f"{args.json_path.stem}.hwp-friendly.rtf"
    docx_path = args.output_dir / f"{args.json_path.stem}.hwp-friendly.docx"
    rtf_path.write_text(render_rtf(doc), encoding="utf-8")
    render_docx(doc, docx_path)
    print(f"RTF: {rtf_path}")
    print(f"DOCX: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
